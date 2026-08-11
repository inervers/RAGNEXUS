import base64
import io

import pytest
from docx import Document

from document_ingest import (
    DocumentIngestError,
    MAX_UPLOAD_BYTES,
    build_preview,
    import_uploaded_document,
    parse_uploaded_document,
)


def _encode_docx(document: Document) -> str:
    stream = io.BytesIO()
    document.save(stream)
    return base64.b64encode(stream.getvalue()).decode("ascii")


def test_preview_truncation_never_changes_import_source() -> None:
    text = "知" * 6001
    encoded = base64.b64encode(text.encode("utf-8")).decode("ascii")

    parsed = parse_uploaded_document("long.txt", encoded)
    preview = build_preview(parsed)

    assert preview == {
        "title": "long",
        "preview": "知" * 5000,
        "full_length": 6001,
        "truncated": True,
    }
    assert parsed.text == text


def test_formal_import_passes_all_parsed_text_to_kb_consumer() -> None:
    text = "完整内容" * 1501
    encoded = base64.b64encode(text.encode("utf-8")).decode("ascii")
    received: list[tuple[str, str]] = []

    def add_document(title: str, content: str) -> dict:
        received.append((title, content))
        return {"message": "added", "chunks": 31}

    result = import_uploaded_document("full.txt", encoded, add_document)

    assert received == [("full", text)]
    assert result == {
        "title": "full",
        "parsed_length": len(text),
        "message": "added",
        "chunks": 31,
    }


def test_upload_rejects_decoded_content_over_limit() -> None:
    encoded = base64.b64encode(b"12345").decode("ascii")

    with pytest.raises(DocumentIngestError, match="超过"):
        parse_uploaded_document("large.txt", encoded, max_bytes=4)


def test_txt_rejects_invalid_utf8_instead_of_replacing_bytes() -> None:
    encoded = base64.b64encode(b"valid\xffinvalid").decode("ascii")

    with pytest.raises(DocumentIngestError, match="UTF-8"):
        parse_uploaded_document("invalid.txt", encoded)


def test_import_rejects_whitespace_only_text_before_kb_mutation() -> None:
    encoded = base64.b64encode(" \n\t".encode("utf-8")).decode("ascii")
    called = False

    def add_document(title: str, content: str) -> dict:
        nonlocal called
        called = True
        return {}

    with pytest.raises(DocumentIngestError, match="有效文本"):
        import_uploaded_document("empty.txt", encoded, add_document)

    assert called is False


def test_empty_pdf_is_a_sanitized_client_validation_error(monkeypatch) -> None:
    import pdf_parser

    monkeypatch.setattr(pdf_parser, "extract_text", lambda raw: {"text": "   "})
    encoded = base64.b64encode(b"fake-pdf").decode("ascii")

    with pytest.raises(DocumentIngestError, match="有效文本") as raised:
        parse_uploaded_document("empty.pdf", encoded)

    assert raised.value.status_code == 400


def test_docx_extracts_paragraphs_and_table_cells_in_body_order() -> None:
    document = Document()
    document.add_paragraph("开头段落")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "表格 A"
    table.cell(0, 1).text = "表格 B"
    table.cell(1, 0).text = "表格 C"
    table.cell(1, 1).text = "表格 D"
    document.add_paragraph("结尾段落")

    parsed = parse_uploaded_document("ordered.docx", _encode_docx(document))

    assert parsed.title == "ordered"
    assert parsed.text == "开头段落\n表格 A\n表格 B\n表格 C\n表格 D\n结尾段落"


def test_docx_parse_error_is_sanitized_for_corrupt_input() -> None:
    encoded = base64.b64encode(b"not-a-docx-package").decode("ascii")

    with pytest.raises(DocumentIngestError, match="DOCX 解析失败") as raised:
        parse_uploaded_document("broken.docx", encoded)

    assert raised.value.status_code == 400


def test_empty_docx_is_rejected_before_kb_mutation() -> None:
    encoded = _encode_docx(Document())
    called = False

    def add_document(title: str, content: str) -> dict:
        nonlocal called
        called = True
        return {}

    with pytest.raises(DocumentIngestError, match="有效文本"):
        import_uploaded_document("empty.docx", encoded, add_document)

    assert called is False


def test_formal_docx_import_passes_complete_extracted_text() -> None:
    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("第二段")
    encoded = _encode_docx(document)
    received: list[tuple[str, str]] = []

    def add_document(title: str, content: str) -> dict:
        received.append((title, content))
        return {"message": "added", "chunks": 1}

    result = import_uploaded_document("report.docx", encoded, add_document)

    assert received == [("report", "第一段\n第二段")]
    assert result["parsed_length"] == len("第一段\n第二段")


def test_default_upload_limit_is_25_mib() -> None:
    assert MAX_UPLOAD_BYTES == 25 * 1024 * 1024
